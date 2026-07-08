"""Extract plain text from an uploaded document, by file type.

Rhinalx ingests the *text* of a source and grounds every decision in a verbatim
span of that text. This module turns an uploaded file's bytes into that text:
PDFs and Word docs are parsed with real libraries; everything else is decoded as
text. Extraction never invents content — if a file has no extractable text (e.g. a
scanned PDF with no text layer), it returns an empty string and the caller refuses.
"""
from __future__ import annotations

import io
from pathlib import Path

# Extensions we can turn into meaningful text. Anything else is still attempted as
# UTF-8/Latin-1 text (many "unknown" files are really plain text).
TEXT_EXTS = {".txt", ".text", ".md", ".markdown", ".csv", ".tsv", ".json", ".log", ".rtf", ""}


def _decode(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts)


def _docx(data: bytes) -> str:
    import docx  # python-docx

    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(filename: str, data: bytes) -> str:
    """Return the plain text of an uploaded file. Empty string means "no text found"."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _pdf(data)
    if ext == ".docx":
        return _docx(data)
    if ext == ".doc":
        # Legacy binary .doc is not supported; ask for .docx/PDF instead.
        raise ValueError("legacy .doc is not supported - please export to .docx or PDF")
    # txt, md, csv, json, logs, or unknown: best-effort text.
    return _decode(data)
